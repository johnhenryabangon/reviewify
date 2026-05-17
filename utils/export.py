"""Export compiled reviewer (multi-lesson) into PDF and DOCX.

Compiled shape:
    {
      "lessons": [
        {"title": str, "source": str, "sections": [section, ...]}, ...
      ]
    }

A section may contain: heading, subheading, summary, bullets,
keywords, formulas, images (filesystem paths or web URLs).
"""
from __future__ import annotations
from pathlib import Path
from typing import List, Dict, Optional


# ============================================================== PDF
def export_pdf(title: str, lessons: List[Dict], out_path: Path,
               static_root: Optional[Path] = None) -> None:
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib.enums import TA_LEFT
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import (
        BaseDocTemplate, Frame, PageTemplate, Paragraph, Spacer,
        ListFlowable, ListItem, Image, KeepTogether, NextPageTemplate,
        PageBreak,
    )

    doc = BaseDocTemplate(
        str(out_path), pagesize=LETTER,
        leftMargin=0.5 * inch, rightMargin=0.5 * inch,
        topMargin=0.6 * inch, bottomMargin=0.6 * inch,
        title=title,
    )
    # Single-column cover frame + two-column content frames
    cover = Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height,
                  id="cover", showBoundary=0)
    gap = 0.25 * inch
    col_w = (doc.width - gap) / 2
    f1 = Frame(doc.leftMargin, doc.bottomMargin, col_w, doc.height,
               id="c1", showBoundary=0)
    f2 = Frame(doc.leftMargin + col_w + gap, doc.bottomMargin, col_w,
               doc.height, id="c2", showBoundary=0)
    doc.addPageTemplates([
        PageTemplate(id="cover", frames=[cover]),
        PageTemplate(id="2col", frames=[f1, f2]),
    ])

    styles = getSampleStyleSheet()
    h_cover = ParagraphStyle("cov", parent=styles["Title"], fontSize=28,
                             spaceAfter=14, textColor=HexColor("#0f172a"))
    h_lesson = ParagraphStyle("hl", parent=styles["Heading1"], fontSize=16,
                              textColor=HexColor("#1e293b"),
                              spaceBefore=4, spaceAfter=8)
    h_sec = ParagraphStyle("hs", parent=styles["Heading2"], fontSize=11,
                           textColor=HexColor("#1f2937"),
                           spaceBefore=8, spaceAfter=3)
    h_sub = ParagraphStyle("hsub", parent=styles["Italic"], fontSize=9,
                           textColor=HexColor("#475569"), spaceAfter=3)
    body = ParagraphStyle("rb", parent=styles["BodyText"], fontSize=9,
                          leading=12, alignment=TA_LEFT)
    bullet = ParagraphStyle("rbu", parent=body, leftIndent=10,
                            bulletIndent=0)
    kw = ParagraphStyle("kw", parent=body, fontSize=8,
                        textColor=HexColor("#4338ca"))
    fm = ParagraphStyle("fm", parent=body, fontName="Courier", fontSize=8.5,
                        textColor=HexColor("#0f766e"))

    story = []
    # Cover
    story.append(Paragraph(_safe(title), h_cover))
    story.append(Paragraph(
        f"Compiled reviewer · {len(lessons)} lesson(s)", body))
    for i, lesson in enumerate(lessons, 1):
        story.append(Spacer(1, 6))
        story.append(Paragraph(
            f"{i}. {_safe(lesson.get('title', 'Lesson'))} "
            f"<font color='#94a3b8'>· {_safe(lesson.get('source',''))}</font>",
            body))
    story.append(NextPageTemplate("2col"))
    story.append(PageBreak())

    for i, lesson in enumerate(lessons, 1):
        story.append(Paragraph(
            f"Lesson {i}: {_safe(lesson.get('title','Lesson'))}", h_lesson))
        for s in lesson.get("sections", []):
            block = []
            block.append(Paragraph(_safe(s.get("heading", "")), h_sec))
            if s.get("subheading"):
                block.append(Paragraph(_safe(s["subheading"]), h_sub))
            if s.get("summary"):
                block.append(Paragraph(_safe(s["summary"]), body))
            if s.get("bullets"):
                items = [ListItem(Paragraph(_safe(b), bullet),
                                  leftIndent=10) for b in s["bullets"]]
                block.append(ListFlowable(items, bulletType="bullet",
                                          start="•", leftIndent=10))
            if s.get("formulas"):
                for f in s["formulas"]:
                    block.append(Paragraph("ƒ  " + _safe(f), fm))
            if s.get("keywords"):
                block.append(Paragraph(
                    "<b>Keywords:</b> " + ", ".join(_safe(k) for k in s["keywords"]),
                    kw))
            for img_url in (s.get("images") or [])[:2]:
                p = _resolve_image(img_url, static_root)
                if not p:
                    continue
                try:
                    img = Image(str(p))
                    img._restrictSize(col_w - 6, 2.4 * 72)
                    block.append(Spacer(1, 3))
                    block.append(img)
                except Exception:
                    pass
            block.append(Spacer(1, 6))
            story.append(KeepTogether(block))
        if i != len(lessons):
            story.append(Spacer(1, 12))
    doc.build(story)


# ============================================================== DOCX
def export_docx(title: str, lessons: List[Dict], out_path: Path,
                static_root: Optional[Path] = None) -> None:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.section import WD_SECTION
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    h = doc.add_heading(title, level=0)
    for run in h.runs:
        run.font.size = Pt(22)
    doc.add_paragraph(f"Compiled reviewer — {len(lessons)} lesson(s)")

    # switch to two-column layout
    section = doc.add_section(WD_SECTION.CONTINUOUS)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    sectPr = section._sectPr
    cols = sectPr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sectPr.append(cols)
    cols.set(qn("w:num"), "2")
    cols.set(qn("w:space"), "360")

    for i, lesson in enumerate(lessons, 1):
        lh = doc.add_heading(
            f"Lesson {i}: {lesson.get('title', 'Lesson')}", level=1)
        for r in lh.runs:
            r.font.size = Pt(14)
            r.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
        for s in lesson.get("sections", []):
            heading = doc.add_heading(s.get("heading", ""), level=2)
            for r in heading.runs:
                r.font.size = Pt(11)
            if s.get("subheading"):
                p = doc.add_paragraph(s["subheading"])
                for r in p.runs:
                    r.italic = True
                    r.font.size = Pt(9)
            if s.get("summary"):
                p = doc.add_paragraph(s["summary"])
                for r in p.runs:
                    r.font.size = Pt(10)
            for b in s.get("bullets", []):
                p = doc.add_paragraph(b, style="List Bullet")
                for r in p.runs:
                    r.font.size = Pt(10)
            for f in s.get("formulas", []):
                p = doc.add_paragraph("ƒ  " + f)
                for r in p.runs:
                    r.font.name = "Consolas"
                    r.font.size = Pt(9)
                    r.font.color.rgb = RGBColor(0x0F, 0x76, 0x6E)
            if s.get("keywords"):
                p = doc.add_paragraph()
                r0 = p.add_run("Keywords: ")
                r0.bold = True
                r0.font.size = Pt(9)
                r1 = p.add_run(", ".join(s["keywords"]))
                r1.font.size = Pt(9)
                r1.font.color.rgb = RGBColor(0x43, 0x38, 0xCA)
            for img_url in (s.get("images") or [])[:2]:
                p = _resolve_image(img_url, static_root)
                if not p:
                    continue
                try:
                    doc.add_picture(str(p), width=Inches(3.0))
                except Exception:
                    pass

    doc.save(str(out_path))


# --------------------------------------------------------------- helpers
def _safe(text: str) -> str:
    return ((text or "")
            .replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;"))


def _resolve_image(url_or_path: str, static_root: Optional[Path]) -> Optional[Path]:
    """Map a /static/... web URL back to a filesystem path."""
    if not url_or_path:
        return None
    p = Path(url_or_path)
    if p.exists():
        return p
    if static_root and url_or_path.startswith("/static/"):
        candidate = static_root / url_or_path[len("/static/"):]
        if candidate.exists():
            return candidate
    return None
